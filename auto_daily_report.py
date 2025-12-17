import os 
import sys
import time
import re
import feedparser
import docx
import requests  # for Telegram
from bs4 import BeautifulSoup
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from collections import defaultdict

# ---------------------------------------------------------------------
# TELEGRAM CONFIG – values will come from environment variables
# (in GitHub Actions: secrets.BOT_TOKEN and secrets.CHAT_ID)
# ---------------------------------------------------------------------
BOT_TOKEN = os.getenv("BOT_TOKEN")
CHAT_ID = os.getenv("CHAT_ID")  # e.g. "-5051605824"

# ---------------------------------------------------------------------
# RSS FEEDS (foreign + Indian) – UNCHANGED
# ---------------------------------------------------------------------

RSS_FEEDS = [
    # Foreign / international (original ones)
    ("BBC", "http://feeds.bbci.co.uk/news/rss.xml"),
    ("CNN", "http://rss.cnn.com/rss/edition.rss"),
    ("Reuters", "http://feeds.reuters.com/reuters/topNews"),
    ("Al Jazeera", "https://www.aljazeera.com/xml/rss/all.xml"),
    ("The Guardian", "https://www.theguardian.com/world/rss"),

    # Indian – with common RSS endpoints (some may need tweaking)
    ("Indian Express", "https://indianexpress.com/section/india/feed/"),
    ("Hindustan Times", "https://www.hindustantimes.com/feeds/rss/india-news/rssfeed.xml"),
    ("The Times of India", "https://timesofindia.indiatimes.com/rssfeedstopstories.cms"),
    ("India Today", "https://www.indiatoday.in/rss/1206578"),
    ("Business Standard", "https://www.business-standard.com/rss/home_page_top_stories.rss"),
    ("First Post", "https://www.firstpost.com/rss/india.xml"),
    ("PIB", "https://pib.gov.in/RssMain.aspx?ModId=6&Lang=1&Regid=3"),
    ("The Hindu", "https://www.thehindu.com/news/national/feeder/default.rss"),
    ("The Telegraph", "https://www.telegraphindia.com/feeds/rss.jsp?id=3"),
    ("The Wire", "https://thewire.in/feed/"),

    # The remaining sources – present in synopsis even if URL is None.
    ("The Tribune", "https://publish.tribuneindia.com/newscategory/top-headlines/feed/"),
    ("Statesman", "https://thestatesmanindia.com/category/news/latest/feed"),
    ("The New Indian Express", "https://www.newindianexpress.com/rss"),
    ("Deccan Herald", "http://www.deccanherald.com/rss-internal/top-stories.rss"),
    ("IANS", None),
    ("PTI", None),
    ("Hands India", None),        # (likely The Hans India – fill URL if you have it)
    ("ANI", None),
    ("Meghalaya Mirror", None),
    ("Nagaland Times", None),
    ("Telengana Today", "https://telanganatoday.com/feed"),
    ("SIASAT Times", "https://www.siasat.com/feed"),
    ("Babu Shahi", None),
    ("Doordarshan", None),
    ("WION", None),
    ("Zee News", "https://zeenews.india.com/rss.html"),
    ("CNBC TV18", None),
    ("Economic Times", "http://economictimes.indiatimes.com/news/economy/rssfeeds/1373380680.cms"),
    ("Financial Express", "https://www.financialexpress.com/section/economy/feed/"),
    ("Mint", "https://www.livemint.com/rss"),
    ("Live Mint", "https://www.livemint.com/rss/news"),
    ("Live Markets", "https://www.livemint.com/rss/markets"),
    ("live Economy/Markets", "http://www.livemint.com/rss/economy_politics"),
    ("Quint", None),
    ("The Print", "https://theprint.in/feed/"),
    ("The Syndication", "http://syndication.financialexpress.com/rss/latest-news.xml")
]

# ---------------------------------------------------------------------
# KEYWORDS – tags / topics you want to track – UNCHANGED
# ---------------------------------------------------------------------

RAW_KEYWORDS = [
    "Ashok Mittal",
    "Ashok Kumar Mittal",
    "Ashok Kumar Mittal Lovely Professional University",
    "lovely professional university",
    "education",
    "Education plus policy",
    "Education plus policy India",
    "Education plus policy center",
    "Education plus policy state",
    "foriegn policy India",
    "trump tariff",
    "pakistan",
    "china",
    "united states america",
    "europe",
    "parliament",
    "winter session",
    "monsoon session",
    "budget session",
    "legislative business",
    "stray dogs",
    "goverment business",
    "minister of education",
    "state government education india",
    "dharmendra pradhan",
    "minister of skilling",
    "development and entreprenurship",
    "Ministry", "Skill", "Development", "Entrepreneurship",
    "upskilling",
    "jayant chaudhury",
    "niti aayog",
    "ministery of women and child development",
    "ministery of health",
    "women and child health",
    "Corporate", "Social", "Responsibility",
    "non-governmental organisation",
    "ministery", "home", "affairs",
    "Foreign Contribution Regulation Act",
    "supreme court",
    "high court",
    "smile foundation",
    "magic bus foundation",
    "pratham education",
    "help age india foundation",
    "santanu mishra",
    "smile on wheels",
    "step smile foundation",
    "United Nations Children Fund",
    "United", "Nations", "Framework", "Convention", "Climate", "Change",
    "United", "Nations", "General", "Assembly",
    "United", "Nations", "High", "Commissioner", "Refugees",
    "United Nations Population Fund",
    "steel",
    "ministery", "steel",
    "green steel",
    "quality check orders",
    "World Trade Organization"
    "steel import",
    "steel export",
    "sd kumaraswamy",
    "steel policy",
    "steel india",
    "steel government",
    "steel global",
    "steel decarbonisation",
    "steel china",
    "issda",
    "Indian Stainless Steel Development Association"
    "isa",
    "International Solar Alliance",
    "wsa",
    "Wilderness Study Area"
    "War Shipping Administration",
    "cbam",
    "Carbon Border Adjustment Mechanism",
    "carbon border tax",
    "steel countervaling duties india",
    "steel tariffs",
    "wto",
    "World Trade Organisation",
    "World Trade Organization",
    "jindal steel",
    "jsw steel",
    "sail",
    "Steel Authority India Limited",
    "tata steel",
    "arcelormittal",
    "policy india",
    "healthcare",
    "jindal south west",
    "Reserve Bank India",
    "rbi",
    "economy india",
    "government india",
    "carbon",
    "climate change",
    "electric vehicles",
    "infrastructure",
    "ai india",
    "ai legal india",
    "ai law india",
    "pollution",
    "meity",
    "indian airforce",
    "indian navy",
    "indian army",
    "military",
    "armed forces",
    "garbage",
    "The Ministry Electronics Information Technology",
    "ministery of information",
    "electronics and technologies",
    "data protection",
    "internet shutdown",
    "data breach",
    "data leak",
    "misinformation",
    "content take down",
    "account blocking",
    "content blocking",
    "cyber security",
    "cyber attack",
    "social media",
    "dpdpa", "act",
    "google ai",
    "microsoft ai",
    "deepseek",
    "youtube",
    "grok",
    "perplexity",
    "chatgpt",
    "open ai",
    "ai tech corporate",
    "ai women",
    "aihealth",
    "indian economy",
    "sebi",
    "Securities Exchange Board India",
    "capital markets",
    "primary markets",
    "gift city",
    "Gujarat International Finance Tec-City",
    "singapore",
    "singapore finance",
    "united", "arab", "emirates", "finance",
    "mauritius finance",
    "alternate investment funds",
    "aif",
    "Alternative Investment Fund",
    "aif sebi",
    "category1 aif",
    "category2 aif",
    "hedge funds",
    "special situation funds",
    "offshore funds",
    "asset reconstruction companies",
    "stressed asset",
    "gift ifsc based funds",
    "nri",
    "nri funds",
    "fdi",
    "Foreign Direct Investment",
    "fpi",
    "Foreign Portfolio Investment",
    "mutual funds",
    "large cap",
    "small cap",
    "digital accessibility funds",
    "fii",
    "bse",
    "nse",
    "hni",
    "nri kyc",
    "family offices",
    "zinnia investments",
    "provizia",
    "safron",
    "pte",
    "quaoar",
    "synergy group",
    "cross border wealth",
    "sme capital",
    "ipo",
    "readiness assessment",
    "market tax",
    "ministery of finance",
    "pre budget consultation",
    "financial stability and development council",
    "ktr",
    "ktr rama rao",
    "Kalvakuntla Taraka Rama Rao",
    "Kalvakuntla Chandrashekar Rao",
    "brs",
    "Building Regularisation Scheme",
    "brs telengana",
    "kcr",
    "harish rao",
    "bjp",
    "congress",
    "trinamool",
    "rahul gandhi",
    "abhishek banerjee",
    "samajvadi party",
    "akilesh yadav",
    "mental health",
    "owaisi",
    "narendra modi",
    "amit shah",
    "nitin gadkari",
    "yogi adityanath",
    "jp nadda",
    "mamata",
    "piyush goyal",
    "shivraj singh chauhan",
    "nitish kumar",
    "chirag paswan",
    "goa",
    "bangladesh",
    "viriato", "fernandes",
    "Odisha education board merger",
    "merger of Class X and Class XII boards",
    "Odisha board merger Class 10 Class 12",
    "Odisha government education board merger",
    "unified education board Odisha",
    "single education board Odisha",
    "Board of Secondary Education Odisha",
    "Council of Higher Secondary Education Odisha",
    "BSE CHSE merger",
    "BSE CHSE integration",
    "Odisha education reform 2025",
    "School and Mass Education Department Odisha",
    "education board merger",
    "state education board integration",
    "education board unification India",
    "govt plans to merge education boards",
    "secondary and higher secondary board merger",
    "Assam education board merger",
    "ASSEB board merger model",
    "karnataka government",
    "wine education course",
    "wine education appreciation and tasting",
    "weat course",
    "wine training bengaluru",
    "karnataka grape and wine board",
    "wine education funding",
    "hospitality training course",
    "Human Rights Day",
    "Legal Education",
    "Human Rights Awareness",
    "Human Rights Law",
    "Role of Legal Education",
    "Justice and Equality",
    "Rule of Law",
    "Human Rights Protection",
    "Kakatiya University",
    "Legal Experts",
    "Legal Literacy",
    "Access to Justice",
    "Social Justice",
    "Constitutional Rights",
    "Human Rights Education in India",
    "Law and Society",
    "Telangana News",
    "Indian Legal System",
    "Higher Education India",
    "unesco_report",
    "right_to_education",
    "global_education_crisis",
    "education_disruptions",
    "education_inequality",
    "out_of_school_children",
    "covid19_education_impact",
    "school_attacks_conflict_zones",
    "education_access",
    "inclusive_education",
    "education_policy",
    "education_recovery",
    "teacher_training",
    "global_learning",
    "human_rights_education",
    "sdg4",
    "education_as_human_right",
    "unesco_recommendations",
    "education_funding",
    "lifelong_learning",
    "dog bites",
    "dog bite prevention",
    "children safety",
    "school safety",
    "student safety",
    "tamil nadu education department",
    "school guidelines",
    "rabies awareness",
    "stray dogs",
    "pet dog vaccination",
    "public health",
    "school health guidelines",
    "animal safety",
    "stray animal management",
    "awareness programs",
    "municipal animal control",
    "tamil nadu news",
    "education and health",
    "higher-education-commission-of-india",
    "heci-bill",
    "heci-bill-2025",
    "higher-education-reform-india",
    "education-policy-india",
    "institutional-autonomy",
    "university-autonomy-india",
    "academic-freedom",
    "ugc-replacement",
    "education-regulatory-reform",
    "parliamentary-debate",
    "cpim-mp",
    "john-brittas",
    "opaque-drafting",
    "federalism-in-education",
    "centralised-education-regulator",
    "higher-education-governance",
    "stakeholder-consultation",
    "national-education-policy-2020",
    "education-news-india",
    "CBSE schools",
    "CBSE Australia",
    "Indian schools in Australia",
    "Union Education Ministry",
    "Dharmendra Pradhan",
    "India Australia education collaboration",
    "Indian diaspora education",
    "Australia India Education and Skills Council",
    "international education policy",
    "CBSE international expansion",
    "Indian curriculum abroad",
    "education diplomacy",
    "foreign education demand",
    "school expansion overseas",
    "India Australia bilateral relations",
    "global Indian education",
    "international school partnerships",
    "karnataka-higher-education",
    "traditional-universities",
    "university-pension-crisis",
    "retired-staff-pensions",
    "public-university-finances",
    "karnataka-universities",
    "higher-education-policy",
    "education-sector-funding",
    "state-government-education",
    "mc-sudhakar",
    "pension-liabilities",
    "university-retirement-benefits",
    "financial-stress-universities",
    "education-minister-statement",
    "legislative-council-karnataka",
    "Bihar schools digital access labs",
    "Bihar education infrastructure",
    "Digital learning Bihar",
    "School technology labs India",
    "Bihar education policy",
    "Government schools digital labs",
    "Education minister Bihar",
    "Smart classrooms Bihar",
    "Digital education India",
    "Education infrastructure news",
    "KGBV",
    "Kasturba Gandhi Balika Vidyalaya",
    "84.2% retention middle school",
    "370 non-functional KGBVs",
    "Education Ministry India",
    "girls education India",
    "UDISE+ data",
    "school dropout rates",
    "middle school retention",
    "secondary school completion",
    "girls' schools",
    "residential schools",
    "education policy",
    "Indian education news",
    "Rajya Sabha education data",
    "school enrollment statistics",
    "functional KGBVs",
    "non functional schools",
    "teacher pupil ratio",
    "Telangana education",
    "Uttar Pradesh KGBV",
    "Andhra Pradesh KGBV",
    "Jharkhand education challenges",
    "ANI News",
    "Jhelum Valley",
    "Kashmir",
    "Pakistan Occupied Kashmir",
    "PoK",
    "Healthcare Crisis",
    "Education Crisis",
    "Public Health",
    "Healthcare Infrastructure",
    "Education System",
    "Political Appeal",
    "Pakistani Politics",
    "Shehbaz Sharif",
    "Pakistan Prime Minister",
    "Abbasi Statement",
    "Regional Crisis",
    "South Asia",
    "Asia News",
    "Governance Failure",
    "Humanitarian Issues",
    "Policy Intervention",
    "UK education policy",
    "Conservative government education",
    "free schools policy UK",
    "free schools failure",
    "UK school reform",
    "education policy analysis UK",
    "academy schools UK",
    "state funded free schools",
    "UK education system criticism",
    "public education accountability",
    "school performance UK",
    "education spending UK",
    "education policy report",
    "UK government education review",
    "education outcomes England",
    "rod paige",
    "rod paige death",
    "rod paige obituary",
    "rod paige education secretary",
    "us education secretary",
    "first black us education secretary",
    "no child left behind act",
    "nclb education reform",
    "education policy united states",
    "federal education policy",
    "education reform",
    "african american leaders",
    "education leadership legacy",
    "education news",
    "demographics education",
    "public education united states",
    "school reform history",
]

DEFAULT_IMAGE_NAMES = {
    "image_1": "image_1",
    "image_2": "image_2",
    "image_3": "image_3",
    "image_4": "image_4",
}

def should_run_now_ist():
    """
    Decide whether to actually run the heavy job now.
    We only want to do real work once per day around 06:00 IST.
    GitHub cron may be delayed, so we check current IST time
    and skip all other hours.
    """
    now_utc = datetime.utcnow()
    ist = now_utc + timedelta(hours=5, minutes=30)
    print("[INFO] Current IST time:", ist.strftime("%Y-%m-%d %H:%M"))

    # Only run when IST hour is 6 (06:00–06:59)
    if ist.hour == 6:
        print("[INFO] Within 06:00–06:59 IST window, proceeding with report.")
        return True

    print("[INFO] Outside 06:00–06:59 IST window, skipping this run.")
    return False
# ---------------------------------------------------------------------
# Utility functions – UNCHANGED
# ---------------------------------------------------------------------

def parse_date_struct(dt_struct):
    try:
        return datetime.fromtimestamp(time.mktime(dt_struct))
    except Exception:
        return None

def first_n_sentences(text, n=3, max_chars=4000):
    sentences = []
    cur = ""
    for ch in text:
        cur += ch
        if ch in ".!?":
            s = cur.strip()
            if s:
                sentences.append(s)
            cur = ""
            if len(sentences) >= n:
                break
    if len(sentences) < n and cur.strip():
        sentences.append(cur.strip())
    result = " ".join(sentences)
    if max_chars is not None:
        return result[:max_chars]
    return result

def get_text_from_html(html_content):
    if not html_content:
        return ""
    soup = BeautifulSoup(html_content, "lxml")
    for tag in soup(["script", "style", "aside", "footer", "nav", "form"]):
        tag.decompose()
    return soup.get_text(separator=" ", strip=True)

def get_matching_tags(item):
    """
    Return a list of RAW_KEYWORDS (tags) that match this item
    based on title / summary / source.

    Matching rules (ALL case-insensitive):
    - Single-word tags:
        * Match as exact tokens only (goa != goal).
    - Multi-word tags:
        * Phrase match with word boundaries.
    - Special case 'act':
        * Only match when text contains a law name like
          'Finance Act', 'Data Protection Act', etc.
    """
    full_text = " ".join([
        item.get("title", "") or "",
        item.get("summary", "") or "",
        item.get("source", "") or "",
    ])

    haystack_lower = full_text.lower()
    tokens = re.findall(r"\w+", haystack_lower)

    matched = []

    for raw in RAW_KEYWORDS:
        kw = (raw or "").strip()
        if not kw:
            continue

        kw_lower = kw.lower()

        # 'act' – law context
        if kw_lower == "act":
            law_pattern = r"\b[A-Z][A-Za-z]*(?:\s+[A-Z][A-Za-z]*)*\s+Act\b"
            if re.search(law_pattern, full_text):
                matched.append(raw)
            continue

        # Multi-word keyword -> phrase match (case-insensitive)
        if " " in kw_lower:
            words = kw_lower.split()
            pattern = r"\b" + r"\s+".join(re.escape(w) for w in words) + r"\b"
            if re.search(pattern, haystack_lower, flags=re.IGNORECASE):
                matched.append(raw)
        else:
            # Single-word keyword -> exact token match, case-insensitive
            if kw_lower in tokens:
                matched.append(raw)

    # dedupe preserving order
    seen = set()
    unique = []
    for t in matched:
        if t not in seen:
            seen.add(t)
            unique.append(t)
    return unique

def fetch_rss_items():
    """
    Fetch items from all configured RSS feeds.
    Tries full content (entry.content) first, then summary/description.
    Uses published_parsed or updated_parsed for date.
    """
    print("[INFO] Starting RSS fetch for all feeds...")
    items = []
    for source_name, feed_url in RSS_FEEDS:
        if not feed_url:
            print(f"[INFO] Skipping source (no URL): {source_name}")
            continue
        try:
            print(f"[INFO] Fetching feed: {source_name} -> {feed_url}")
            d = feedparser.parse(feed_url)
            print(f"[INFO] Parsed feed: {source_name}, entries: {len(getattr(d, 'entries', []))}")
            for entry in d.entries:
                published = None
                if hasattr(entry, "published_parsed") and entry.published_parsed:
                    published = parse_date_struct(entry.published_parsed)
                elif hasattr(entry, "updated_parsed") and entry.updated_parsed:
                    published = parse_date_struct(entry.updated_parsed)

                link = entry.get("link", "")
                title = entry.get("title", "No title")

                # Prefer full "content" over "summary"
                html_content = ""
                if hasattr(entry, "content"):
                    try:
                        if entry.content and len(entry.content) > 0:
                            html_content = entry.content[0].get("value", "") or ""
                    except Exception:
                        html_content = ""
                if not html_content:
                    html_content = entry.get("summary", "") or entry.get("description", "") or ""

                summary_text = get_text_from_html(html_content)

                items.append({
                    "source": source_name,
                    "title": title,
                    "link": link,
                    "summary": summary_text,
                    "published": published,
                })
        except Exception as e:
            print(f"[ERROR] Failed to parse {feed_url} ({source_name}): {e}", file=sys.stderr)
    print(f"[INFO] Finished fetching RSS items. Total items: {len(items)}")
    return items

def add_hyperlink(paragraph, url, text, color="0000FF", underline=True, bold=True):
    """
    Insert a clickable hyperlink into a python-docx paragraph.
    """
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    c = OxmlElement("w:color")
    c.set(qn("w:val"), color)
    rPr.append(c)

    if underline:
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rPr.append(u)

    if bold:
        b = OxmlElement("w:b")
        b.set(qn("w:val"), "true")
        rPr.append(b)

    new_run.append(rPr)

    text_elem = OxmlElement("w:t")
    text_elem.text = text
    new_run.append(text_elem)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink

# ---------------------------------------------------------------------
# DOCX builder – with source-wise synopsis + detailed section – UNCHANGED
# (uses shorter 3-sentence, ~800-char summaries via first_n_sentences)
# ---------------------------------------------------------------------

def build_docx(detailed_items,
               per_source_all,
               per_source_relevant,
               input_datetime,
               image_paths,
               output_path,
               since_dt,
               until_dt):
    print("[INFO] Building DOCX report...")
    doc = Document()

    # PAGE 1: Cover image
    if image_paths.get("image_1") and os.path.exists(image_paths["image_1"]):
        print(f"[INFO] Adding cover image: {image_paths['image_1']}")
        p = doc.add_paragraph()
        run = p.add_run()
        run.add_picture(image_paths["image_1"], width=Inches(6.5), height=Inches(9))
    else:
        print("[INFO] Cover image (image_1) not found; skipping.")

    p_report = doc.add_paragraph("Report")
    p_report.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Date under Report, left justified
    doc.add_paragraph("")
    doc.add_paragraph("")
    p_date = doc.add_paragraph(input_datetime.strftime("%d.%m.%Y"))
    p_date.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    p_date.runs[0].font.name = "Times New Roman"
    p_date.runs[0].font.size = Pt(28)

    doc.add_page_break()

    # PAGE 2 onward: header/footer images
    sec = doc.sections[-1]
    header = sec.header
    footer = sec.footer

    if image_paths.get("image_2") and os.path.exists(image_paths["image_2"]):
        print(f"[INFO] Adding header image: {image_paths['image_2']}")
        p_head = header.add_paragraph()
        run_head = p_head.add_run()
        run_head.add_picture(image_paths["image_2"], width=Inches(6.5))
    else:
        print("[INFO] Header image (image_2) not found; skipping.")

    if image_paths.get("image_3") and os.path.exists(image_paths["image_3"]):
        print(f"[INFO] Adding footer image: {image_paths['image_3']}")
        p_footer = footer.add_paragraph()
        run_footer = p_footer.add_run()
        run_footer.add_picture(image_paths["image_3"], width=Inches(6.5))
    else:
        print("[INFO] Footer image (image_3) not found; skipping.")

    # -----------------------------------------------------------------
    # Source-wise Synopsis – every feed is mentioned
    # -----------------------------------------------------------------
    print("[INFO] Adding Source-wise Synopsis section...")
    syn_heading = doc.add_paragraph("Source-wise Synopsis")
    syn_heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    syn_run = syn_heading.runs[0]
    syn_run.bold = True
    syn_run.font.size = Pt(16)

    period_para = doc.add_paragraph(
        f"Coverage window: {since_dt.strftime('%d-%m-%Y %H:%M')} to {until_dt.strftime('%d-%m-%Y %H:%M')}"
    )
    period_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    for source_name, feed_url in RSS_FEEDS:
        p_source = doc.add_paragraph()
        name_run = p_source.add_run(source_name)
        name_run.bold = True

        if not feed_url:
            p_source.add_run(" – feed URL not configured; no articles pulled.")
            continue

        all_items = per_source_all.get(source_name, [])
        rel_items = per_source_relevant.get(source_name, [])

        if not all_items:
            p_source.add_run(" – no articles in this period.")
            continue

        p_source.add_run(
            f" – {len(all_items)} articles in this period; "
            f"{len(rel_items)} matched your keyword set."
        )

        # Brief synopsis:
        # If there are relevant (keyword-matched) items, show ALL of them.
        # If there are none, show up to 3 general articles for context.
        if rel_items:
            list_items = rel_items
        else:
            list_items = all_items[:3]

        for it in list_items:
            bp = doc.add_paragraph()
            bp.paragraph_format.left_indent = Inches(0.25)
            bullet_run = bp.add_run("• ")
            bullet_run.bold = True
            add_hyperlink(bp, it["link"], it["title"])
            snippet = first_n_sentences(it.get("summary", ""), n=1, max_chars=300)
            if snippet:
                bp.add_run(" – " + snippet)

            # Tags in synopsis: bold, uppercase, red
            tags = it.get("matched_tags", [])
            if tags:
                bp.add_run(" [")
                label_run = bp.add_run("Tags: ")
                label_run.bold = True
                for idx, tag in enumerate(tags):
                    if idx > 0:
                        bp.add_run(", ")
                    r_tag = bp.add_run(tag.upper())
                    r_tag.bold = True
                    r_tag.font.color.rgb = RGBColor(255, 0, 0)
                bp.add_run("]")

    doc.add_page_break()

    # -----------------------------------------------------------------
    # Detailed Articles section – keyword-matched, SHORTER summaries
    # -----------------------------------------------------------------
    print("[INFO] Adding Detailed Articles section...")
    detail_heading = doc.add_paragraph("Detailed Articles (Keyword-Matched)")
    detail_heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    d_run = detail_heading.runs[0]
    d_run.bold = True
    d_run.font.size = Pt(16)

    if not detailed_items:
        doc.add_paragraph("No articles matched your keyword list in this period.")
    else:
        for it in detailed_items:
            # Headline with hyperlink
            p_headline = doc.add_paragraph()
            add_hyperlink(p_headline, it["link"], it["title"], color="0000FF")

            # Meta info
            p_meta = doc.add_paragraph()
            r1 = p_meta.add_run("Source: ")
            r1.bold = True
            p_meta.add_run(it.get("source", "Unknown"))
            pub = it.get("published")
            if pub:
                p_meta.add_run("   |   ")
                r2 = p_meta.add_run("Published: ")
                r2.bold = True
                p_meta.add_run(pub.strftime("%d-%m-%Y %H:%M"))

            # Raw URL line
            doc.add_paragraph(f"Link: {it['link']}")

            # Tags line – bold, uppercase, red for keywords
            tags = it.get("matched_tags", [])
            if tags:
                p_tags = doc.add_paragraph()
                label_run = p_tags.add_run("Tags: ")
                label_run.bold = True
                for idx, tag in enumerate(tags):
                    if idx > 0:
                        p_tags.add_run(", ")
                    r_tag = p_tags.add_run(tag.upper())
                    r_tag.bold = True
                    r_tag.font.color.rgb = RGBColor(255, 0, 0)

            # Shorter detailed summary – up to 3 sentences, 800 chars
            full_summary = it.get("summary", "") or ""
            detailed_text = first_n_sentences(full_summary, n=3, max_chars=800)
            if detailed_text:
                doc.add_paragraph(detailed_text)

            # Blank line between articles
            doc.add_paragraph("")

    # LAST PAGE: only image_4
    doc.add_page_break()
    if image_paths.get("image_4") and os.path.exists(image_paths["image_4"]):
        print(f"[INFO] Adding final page image: {image_paths['image_4']}")
        p_end = doc.add_paragraph()
        p_end.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run_end = p_end.add_run()
        run_end.add_picture(image_paths["image_4"], width=Inches(6.5))
    else:
        print("[INFO] Final page image (image_4) not found; skipping.")

    print(f"[INFO] Saving DOCX report to: {output_path}")
    doc.save(output_path)
    print("[INFO] DOCX report build completed.")

# ---------------------------------------------------------------------
# Main report logic – UNCHANGED, but called from auto main()
# ---------------------------------------------------------------------

def run_report(input_dt, image_paths, output_fullpath):
    since = input_dt - timedelta(hours=24)
    until = input_dt

    print("[INFO] ----------------------------------------------------")
    print(f"[INFO] Starting report generation for window:")
    print(f"[INFO]   From: {since}")
    print(f"[INFO]   To  : {until}")
    print("[INFO] ----------------------------------------------------")

    items = fetch_rss_items()

    print("[INFO] Filtering items within time window and matching tags...")
    in_window = []
    for it in items:
        pub = it.get("published")
        if pub and since <= pub <= until:
            tags = get_matching_tags(it)
            it["matched_tags"] = tags
            in_window.append(it)

    print(f"[INFO] Items in time window: {len(in_window)}")

    per_source_all = defaultdict(list)
    for it in in_window:
        per_source_all[it["source"]].append(it)

    detailed_items = []
    per_source_relevant = defaultdict(list)
    for it in in_window:
        tags = it.get("matched_tags", [])
        if tags:
            detailed_items.append(it)
            per_source_relevant[it["source"]].append(it)

    print(f"[INFO] Items with at least one matched tag: {len(detailed_items)}")
    detailed_items.sort(key=lambda x: x.get("published") or datetime.min, reverse=True)

    build_docx(
        detailed_items,
        per_source_all,
        per_source_relevant,
        input_datetime=input_dt,
        image_paths=image_paths,
        output_path=output_fullpath,
        since_dt=since,
        until_dt=until,
    )

    print("[INFO] Report generation complete.")

# ---------------------------------------------------------------------
# Image auto-detection – based on your existing find_images, but as a
# top-level helper for automation (no Tkinter).
# ---------------------------------------------------------------------

def find_images(out_folder):
    """
    Auto-detect images for image_1..image_4 in several locations
    and with multiple possible extensions.
    """
    script_dir = os.path.dirname(os.path.abspath(sys.argv[0]))
    cwd = os.getcwd()

    search_dirs = []
    for d in [script_dir, out_folder, cwd]:
        if d and d not in search_dirs:
            search_dirs.append(d)

    exts = [".jpeg", ".jpg", ".png", ".gif", ".bmp"]

    image_paths = {}
    print("[INFO] ----------------------------------------------------")
    print("[INFO] Searching for images in these folders (in order):")
    for d in search_dirs:
        print(f"[INFO]   {d}")
    print("[INFO] ----------------------------------------------------")

    for key, base_name in DEFAULT_IMAGE_NAMES.items():
        found = None
        for d in search_dirs:
            for ext in exts:
                candidate = os.path.join(d, base_name + ext)
                if os.path.exists(candidate):
                    found = candidate
                    break
            if found:
                break
        if found:
            image_paths[key] = found
            print(f"[INFO] Found {key}: {found}")
        else:
            image_paths[key] = None
            print(f"[INFO] {key} not found in any search folder; skipping.")
    return image_paths

# ---------------------------------------------------------------------
# Telegram sending
# ---------------------------------------------------------------------

def send_report_to_telegram(file_path):
    if not BOT_TOKEN or not CHAT_ID:
        raise RuntimeError("BOT_TOKEN or CHAT_ID not set in environment.")
    url = f"https://api.telegram.org/bot{BOT_TOKEN}/sendDocument"
    with open(file_path, "rb") as f:
        files = {"document": f}
        data = {
            "chat_id": CHAT_ID,
            "caption": (
                "Good Morning, Kabuze, today's Mittal News Report file "
                "is ready for your review."
            ),
        }
        resp = requests.post(url, data=data, files=files)
    print("[INFO] Telegram response:", resp.status_code, resp.text)
    resp.raise_for_status()

# ---------------------------------------------------------------------
# AUTOMATION MAIN – no UI, suitable for GitHub Actions
# ---------------------------------------------------------------------

def main_auto():
    # Use IST (UTC+5:30) for the report timestamp
    now_utc = datetime.utcnow()
    now_ist = now_utc + timedelta(hours=5, minutes=30)

    out_folder = os.getcwd()
    output_fullpath = os.path.join(
        out_folder,
        f"Daily Report - Dr. Mittal - {now_ist.strftime('%d.%m.%Y')}.docx"
    )

    print("[INFO] ----------------------------------------------------")
    print("[INFO] Auto report generation started (no UI).")
    print(f"[INFO] UTC time: {now_utc.strftime('%Y-%m-%d %H:%M:%S')}")
    print(f"[INFO] IST time: {now_ist.strftime('%Y-%m-%d %H:%M:%S')}")
    print(f"[INFO] Output folder: {out_folder}")
    print(f"[INFO] Output file  : {output_fullpath}")
    print("[INFO] ----------------------------------------------------")

    image_paths = find_images(out_folder)

    run_report(now_ist, image_paths, output_fullpath)

    print("[INFO] Sending report to Telegram group...")
    send_report_to_telegram(output_fullpath)
    print("[INFO] All done.")

if __name__ == "__main__":
    # GitHub Actions cron is adjusted so this effectively runs around 06:00 IST.
    main_auto()

